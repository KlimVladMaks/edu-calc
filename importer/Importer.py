import json
import datetime
from docx import Document
from docx.shared import Pt
from openpyxl import Workbook
from importer.StageRange import StageRange


class Importer:
    """
    Класс для импорта данных об учебных группах из JSON файла (database.json)
    и экспорта их в форматы DOCX и XLSX.
    """
    def __init__(self, json_file_path):
        self.json_file_path = json_file_path
        self.data = self.load_data()
        self.calendars = {cal["name"]: cal for cal in self.data.get("calendars", [])}
        self.groups = self.data.get("groups", [])
        self.programs = {prog["name"]: prog for prog in self.data.get("programs", [])}

    def load_data(self):
        """
        Загружает данные из JSON файла.
        """
        try:
            with open(self.json_file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
                if isinstance(data, dict):
                    return data
                else:
                    print("Ошибка: JSON-файл должен содержать объект с ключами 'calendars', 'edu_types', 'edu_stages', 'programs' и 'groups'.")
                    return {}
        except FileNotFoundError:
            print(f"Файл {self.json_file_path} не найден.")
            return {}
        except json.JSONDecodeError as e:
            print("Ошибка декодирования JSON:", e)
            return {}

    def is_working_day(self, date_obj, calendar):
        """
        Возвращаем True, если date_obj — рабочий день:
         - Не в days_off_list
         - Если "Пятидневный" => суббота(5)/воскресенье(6) = выходные
         - Если "Шестидневный" => только воскресенье(6) = выходной
        """
        days_off = {
            datetime.datetime.strptime(d, "%Y-%m-%d").date()
            for d in calendar.get("days_off_list", [])
        }
        if date_obj in days_off:
            return False

        name_lower = calendar["name"].lower()
        weekday = date_obj.weekday()  # Пн=0, ..., Вс=6

        if "пятиднев" in name_lower:
            if weekday in (5, 6):
                return False
        elif "шестиднев" in name_lower:
            if weekday == 6:
                return False

        return True

    def compute_stages(self, start_date, calendar_name, program_name):
        """
        Возвращает список StageRange (один на каждый этап),
        двигаясь по дням от start_date, пропуская выходные (is_working_day=False).
        """
        calendar = self.calendars.get(calendar_name)
        program = self.programs.get(program_name)
        if not calendar or not program:
            return []

        stages_info = program.get("stages", [])
        current_date = start_date
        results = []

        for (stage_name, stage_duration) in stages_info:
            while not self.is_working_day(current_date, calendar):
                current_date += datetime.timedelta(days=1)

            stage_start = current_date

            needed = stage_duration
            while needed > 0:
                if self.is_working_day(current_date, calendar):
                    needed -= 1
                current_date += datetime.timedelta(days=1)

            stage_end = current_date - datetime.timedelta(days=1)

            results.append(StageRange(stage_name, stage_start, stage_end))

        return results

    def parse_stages_into_columns(self, stages_list):
        """
        Возвращает словарь:
        {
          "theory_start": date or None,
          "theory_end": date or None,
          "practice_start": date or None,
          "practice_end": date or None,
          "prob_trip": [date or (start,end), ...],
          "consult": [...],
          "exam": [...]
        }
        """
        result = {
            "theory_start": None,
            "theory_end": None,
            "practice_start": None,
            "practice_end": None,
            "prob_trip": [],
            "consult": [],
            "exam": []
        }
        for sr in stages_list:
            name = sr.stage_name.lower()
            if "теория" in name:
                if result["theory_start"] is None:
                    result["theory_start"] = sr.start_date
                result["theory_end"] = sr.end_date

            elif "практика" in name:
                if result["practice_start"] is None:
                    result["practice_start"] = sr.start_date
                result["practice_end"] = sr.end_date

            elif "пробная поездка" in name:
                if sr.is_single_day():
                    result["prob_trip"].append(sr.start_date)
                else:
                    result["prob_trip"].append((sr.start_date, sr.end_date))

            elif "консультация" in name:
                if sr.is_single_day():
                    result["consult"].append(sr.start_date)
                else:
                    result["consult"].append((sr.start_date, sr.end_date))

            elif "экзамен" in name:
                if sr.is_single_day():
                    result["exam"].append(sr.start_date)
                else:
                    result["exam"].append((sr.start_date, sr.end_date))
        return result

    def format_date(self, date_obj):
        """Вернём 'дд.мм.гггг' или пусто."""
        if not date_obj:
            return ""
        return date_obj.strftime("%d.%m.%Y")

    def format_date_or_range(self, item):
        """Если item - дата => дд.мм.гггг, если (start,end) => 'дд.мм.гггг – дд.мм.гггг'."""
        if isinstance(item, datetime.date):
            return item.strftime("%d.%m.%Y")
        elif isinstance(item, tuple) and len(item) == 2:
            s = item[0].strftime("%d.%m.%Y")
            e = item[1].strftime("%d.%m.%Y")
            if s == e:
                return s
            else:
                return f"{s} – {e}"
        return ""

    def dash_if_empty(self, text):
        """
        Если text пустая строка, вернуть '-'
        Иначе вернуть text.
        """
        return text if text else "-"

    def compute_row_data(self, group):
        """
        Возвращает кортеж (10 колонок) или None, если не удалось.
        1) Название программы
        2) Вид обучения
        3) Номер группы
        4) Начало теоретического обучения
        5) Конец теоретического обучения
        6) Начало производственной практики
        7) Конец производственной практики
        8) Дата пробной поездки (через запятую)
        9) Консультация (через запятую)
        10) Экзамен (через запятую)
        """
        program_name = group.get("program", "")
        edu_type = group.get("edu_type", "")
        group_name = group.get("name", "")
        calendar_name = group.get("calendar", "")
        start_date_str = group.get("start_date", "")

        if not program_name or not edu_type or not group_name or not start_date_str or not calendar_name:
            print(f"Ошибка: не хватает данных у группы '{group_name}'. Пропускаем.")
            return None

        try:
            start_date = datetime.datetime.strptime(start_date_str, "%Y-%m-%d").date()
        except ValueError:
            print(f"Ошибка: неверная дата '{start_date_str}' в группе '{group_name}'. Пропускаем.")
            return None

        stages_list = self.compute_stages(start_date, calendar_name, program_name)
        if not stages_list:
            print(f"У группы '{group_name}' нет этапов или ошибка в программе/календаре.")
            return None

        col = self.parse_stages_into_columns(stages_list)

        theory_start = self.format_date(col["theory_start"])
        theory_end   = self.format_date(col["theory_end"])

        practice_start = self.format_date(col["practice_start"])
        practice_end   = self.format_date(col["practice_end"])

        prob_list = [self.format_date_or_range(x) for x in col["prob_trip"]]
        prob_str = ", ".join([d for d in prob_list if d])

        consult_list = [self.format_date_or_range(x) for x in col["consult"]]
        consult_str  = ", ".join([d for d in consult_list if d])

        exam_list = [self.format_date_or_range(x) for x in col["exam"]]
        exam_str  = ", ".join([d for d in exam_list if d])

        return (
            program_name,        # 1
            edu_type,            # 2
            group_name,          # 3
            theory_start,        # 4
            theory_end,          # 5
            practice_start,      # 6
            practice_end,        # 7
            prob_str,            # 8
            consult_str,         # 9
            exam_str             # 10
        )

    def filter_groups(self, selected_groups=None):
        if not selected_groups:
            return self.groups
        return [g for g in self.groups if g["name"] in selected_groups]

    def export_docx(self, docx_file_path, selected_groups=None):
        """
        Одна таблица для всех (или выбранных) групп.
        Колонки (10 штук), как на скриншоте:
          1) Название программы
          2) Вид обучения
          3) Номер группы
          4) Начало теоретического обучения
          5) Конец теоретического обучения
          6) Начало производственной практики
          7) Конец производственной практики
          8) Дата пробной поездки
          9) Консультация
          10) Экзамен
        """
        doc = Document()
        doc.add_heading("Учебные группы (единая таблица)", level=1)

        section = doc.sections[-1]
        from docx.enum.section import WD_ORIENT
        section.orientation = WD_ORIENT.LANDSCAPE
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height

        headers = [
            "Название программы",
            "Вид обучения (подготовка/ переподготовка/ повышение квалификации)",
            "Номер группы",
            "Начало теоретического обучения",
            "Конец теоретического обучения",
            "Начало производственной практики",
            "Конец производственной практики",
            "Дата пробной поездки",
            "Консультация",
            "Экзамен"
        ]

        table = doc.add_table(rows=1, cols=len(headers))
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header

        groups_to_export = self.filter_groups(selected_groups)

        for group in groups_to_export:
            row_data = self.compute_row_data(group)
            if row_data is None:
                continue
            row_cells = table.add_row().cells
            for i, val in enumerate(row_data):
                row_cells[i].text = val

        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)

        try:
            doc.save(docx_file_path)
            print(f"Таблица успешно экспортирована в DOCX: {docx_file_path}")
        except Exception as e:
            print("Ошибка сохранения DOCX:", e)

    def export_excel(self, xlsx_file_path, selected_groups=None):
        """
        Аналогичная таблица (10 колонок) в Excel.
        """
        wb = Workbook()
        ws = wb.active
        ws.title = "Учебные группы"

        headers = [
            "Название программы",
            "Вид обучения (подготовка/ переподготовка/ повышение квалификации)",
            "Номер группы",
            "Начало теоретического обучения",
            "Конец теоретического обучения",
            "Начало производственной практики",
            "Конец производственной практики",
            "Дата пробной поездки",
            "Консультация",
            "Экзамен"
        ]
        ws.append(headers)

        groups_to_export = self.filter_groups(selected_groups)

        for group in groups_to_export:
            row_data = self.compute_row_data(group)
            if row_data is None:
                continue
            ws.append(row_data)

        try:
            wb.save(xlsx_file_path)
            print(f"Таблица успешно экспортирована в Excel: {xlsx_file_path}")
        except Exception as e:
            print("Ошибка сохранения Excel:", e)

    def format_theory_practice_range(self, stage_name, start_date, end_date, practice=False):
        """
        Формирует строку для этапов "Теория" или "Практика".
        Если входные параметры не являются датами, пытается их преобразовать.
        Если этап занимает один день, выводит только одну дату.
        Если practice=True, используется формат "с dd.mm.yyyy по dd.mm.yyyy".
        """
        if not isinstance(start_date, datetime.date):
            try:
                start_date = datetime.datetime.strptime(start_date, "%d.%m.%Y").date()
            except Exception:
                return f"{stage_name}: -"
        if not isinstance(end_date, datetime.date):
            try:
                end_date = datetime.datetime.strptime(end_date, "%d.%m.%Y").date()
            except Exception:
                return f"{stage_name}: -"
        start_str = start_date.strftime("%d.%m.%Y")
        end_str = end_date.strftime("%d.%m.%Y")
        if start_date == end_date:
            if practice:
                return f"{stage_name} с {start_str} по {end_str}."
            else:
                return f"{stage_name} {start_str}."
        else:
            if practice:
                return f"{stage_name} с {start_str} по {end_str}."
            else:
                return f"{stage_name} {start_str} – {end_str}."

    def format_multi_dates(self, dates_list, title):
        """
        Формирует строку для этапов, которые могут повторяться (например, Консультации, Экзамены).
        Если дат нет, выводит прочерк.
        Пример: "Консультации: 12.05.2025, 22.08.2025, 26.08.2025."
        """
        items = []
        for item in dates_list:
            if isinstance(item, datetime.date):
                items.append(item.strftime("%d.%m.%Y"))
            elif isinstance(item, tuple) and len(item) == 2:
                s = item[0].strftime("%d.%m.%Y")
                e = item[1].strftime("%d.%m.%Y")
                if s == e:
                    items.append(s)
                else:
                    items.append(f"{s} – {e}")
        joined = ", ".join(items)
        return f"{title}: {joined}." if joined else f"{title}: -."

    def export_group_details_docx(self, docx_file_path, selected_groups=None):
        """
        Выгружает для каждой группы (одной или нескольких) текстовый блок вида:

            Группа М-248
            Теория 24.02.2025 - 07.05.2025., 21.08.2025.
            Практика с 14.05.2025 по 20.08.2025.
            Консультации: 12.05.2025, 22.08.2025, 26.08.2025.
            Экзамены 13.05.2025, 25.08.2025, 27.08.2025.

        Если группа не имеет этапа, строка не выводится или выводится без дат (как захотите).
        """
        doc = Document()
        doc.add_heading("Детальный отчёт по группам", level=1)

        groups_to_export = self.filter_groups(selected_groups)

        for group in groups_to_export:
            group_name = group.get("name", "")
            program_name = group.get("program", "")
            edu_type = group.get("edu_type", "")
            calendar_name = group.get("calendar", "")
            start_date_str = group.get("start_date", "")

            doc.add_heading(f"Группа {group_name}", level=2)

            try:
                start_date = datetime.datetime.strptime(start_date_str, "%Y-%m-%d").date()
            except ValueError:
                doc.add_paragraph(f"  Ошибка: неверная дата начала '{start_date_str}'.")
                continue

            stages_list = self.compute_stages(start_date, calendar_name, program_name)
            if not stages_list:
                doc.add_paragraph("  Нет данных об этапах или ошибка в программе/календаре.")
                continue

            columns = self.parse_stages_into_columns(stages_list)

            if columns["theory_start"] is not None:
                theory_line = self.format_theory_practice_range(
                    "Теория",
                    columns["theory_start"],
                    columns["theory_end"]
                )
                doc.add_paragraph(theory_line)

            if columns["practice_start"] is not None:
                practice_line = self.format_theory_practice_range(
                    "Практика",
                    columns["practice_start"],
                    columns["practice_end"],
                    practice=True
                )
                doc.add_paragraph(practice_line)

            if columns["prob_trip"]:
                trip_line = self.format_multi_dates(columns["prob_trip"], "Пробная поездка")
                doc.add_paragraph(trip_line)

            if columns["consult"]:
                consult_line = self.format_multi_dates(columns["consult"], "Консультации")
                doc.add_paragraph(consult_line)

            if columns["exam"]:
                exam_line = self.format_multi_dates(columns["exam"], "Экзамены")
                doc.add_paragraph(exam_line)

            doc.add_paragraph()

        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)
        doc.save(docx_file_path)
        print(f"Детальный отчёт по группам экспортирован в {docx_file_path}")
